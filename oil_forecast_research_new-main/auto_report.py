from pathlib import Path
import os
import sys
import numpy as np
import pandas as pd
import torch
from docx import Document
from docx.shared import Pt, Inches

# Add current dir to path so src.model.model can be imported
sys.path.append(str(Path(__file__).resolve().parent))
from src.model.model import GUMNet

BASE_DIR = Path(__file__).resolve().parent
DEFAULT_CKPT = BASE_DIR / "checkpoints" / "gumnet_ckpt.pt"
BUILTIN_DATA = BASE_DIR / "data" / "processed" / "clean_data_exo_ver1.csv"
QUANTILE_LABELS = ["p10", "p50", "p90"]
DATASETS_DIR = BASE_DIR.parent / "datasets"
OUTPUT_WORD = BASE_DIR / "Bao_cao_Du_bao_v3.docx"

# ─────────────────────── Helpers copied from app.py ─────────────────────────
def next_business_days(last_date: pd.Timestamp, n: int):
    days = []
    d = pd.Timestamp(last_date)
    while len(days) < n:
        d += pd.Timedelta(days=1)
        if d.weekday() < 5:
            days.append(d)
    return days

def read_df_path(file_path, date_col: str) -> pd.DataFrame:
    file_path = Path(file_path)
    name = file_path.name.lower()
    if name.endswith((".xlsx", ".xls")):
        df = pd.read_excel(file_path)
    else:
        df = pd.read_csv(file_path)

    df[date_col] = pd.to_datetime(df[date_col], errors="coerce")
    df = df.dropna(subset=[date_col]).sort_values(date_col).reset_index(drop=True)

    for c in df.columns:
        if c != date_col:
            df[c] = pd.to_numeric(df[c], errors="coerce")

    num_cols = [c for c in df.columns if c != date_col]
    df[num_cols] = df[num_cols].interpolate(method="linear").bfill().ffill()
    return df

def load_builtin_df(date_col: str) -> pd.DataFrame:
    df = pd.read_csv(BUILTIN_DATA)
    date_candidates = [c for c in df.columns if "ng" in c.lower() or "date" in c.lower() or "day" in c.lower()]
    if date_candidates and date_candidates[0] != date_col:
        df = df.rename(columns={date_candidates[0]: date_col})
    df[date_col] = pd.to_datetime(df[date_col], errors="coerce")
    df = df.dropna(subset=[date_col]).sort_values(date_col).reset_index(drop=True)
    for c in df.columns:
        if c != date_col:
            df[c] = pd.to_numeric(df[c], errors="coerce")
    num_cols = [c for c in df.columns if c != date_col]
    df[num_cols] = df[num_cols].interpolate(method="linear").bfill().ffill()
    return df

def merge_missing_exo(user_df: pd.DataFrame, date_col: str, required_cols: list) -> tuple[pd.DataFrame, list]:
    missing = [c for c in required_cols if c not in user_df.columns]
    if not missing:
        return user_df, []
    if not BUILTIN_DATA.exists():
        return user_df, []
    builtin = load_builtin_df(date_col)
    exo_available = [c for c in missing if c in builtin.columns]
    if not exo_available:
        return user_df, []
    exo_df = builtin[[date_col] + exo_available].copy()
    merged = pd.merge(user_df, exo_df, on=date_col, how="left")
    merged[exo_available] = merged[exo_available].ffill().bfill()
    for col in exo_available:
        if merged[col].isna().any():
            last_val = builtin[col].iloc[-1]
            merged[col] = merged[col].fillna(last_val)
    return merged, exo_available

def load_model(ckpt_path: str):
    device = "cuda" if torch.cuda.is_available() else "cpu"
    ckpt = torch.load(ckpt_path, map_location=device, weights_only=False)

    model = GUMNet(
        seq_len=ckpt["seq_len"],
        input_dim=ckpt["input_dim"],
        output_dim=ckpt["output_dim"],
        horizon=ckpt["horizon"],
        num_quantiles=ckpt["num_quantiles"],
    ).to(device)
    model.load_state_dict(ckpt["model_state_dict"])
    model.eval()
    return model, ckpt, device

def predict_future(model, ckpt, df: pd.DataFrame, device: str):
    feature_cols  = ckpt["feature_cols"]
    target_cols   = ckpt["target_cols"]
    seq_len       = ckpt["seq_len"]
    horizon       = ckpt["horizon"]
    date_col      = ckpt["date_col"]

    X_all = ckpt["feature_scaler"].transform(df[feature_cols].values)
    x_last = torch.tensor(X_all[-seq_len:], dtype=torch.float32).unsqueeze(0).to(device)

    with torch.no_grad():
        pred_scaled, weights = model(x_last)

    pred_np = pred_scaled.cpu().numpy()[0]
    pred_np = np.sort(pred_np, axis=-1)  # Đảm bảo p10 <= p50 <= p90
    results = {}
    for qi, ql in enumerate(QUANTILE_LABELS):
        p = ckpt["target_scaler"].inverse_transform(pred_np[:, :, qi])
        results[ql] = pd.DataFrame(p, columns=target_cols)

    future_dates = next_business_days(pd.to_datetime(df[date_col].iloc[-1]), horizon)
    for ql in QUANTILE_LABELS:
        results[ql].insert(0, date_col, future_dates)

    return results, weights.cpu().numpy()

def backtest(model, ckpt, df: pd.DataFrame, device: str, n_samples: int = 200):
    feature_cols  = ckpt["feature_cols"]
    target_cols   = ckpt["target_cols"]
    seq_len       = ckpt["seq_len"]
    date_col      = ckpt["date_col"]

    X_all   = ckpt["feature_scaler"].transform(df[feature_cols].values)
    dates   = df[date_col].values
    actuals = df[target_cols].values

    start_i = max(seq_len, len(df) - seq_len - n_samples)
    preds_p10, preds_p50, preds_p90 = [], [], []
    actual_list, date_list = [], []

    for i in range(start_i, len(df) - 1):
        x = torch.tensor(X_all[i - seq_len: i], dtype=torch.float32).unsqueeze(0).to(device)
        with torch.no_grad():
            pred_scaled, _ = model(x)
        p_np = pred_scaled.cpu().numpy()[0, 0, :, :]
        p_np = np.sort(p_np, axis=-1)  # Đảm bảo p10 <= p50 <= p90
        def inv(arr):
            return ckpt["target_scaler"].inverse_transform(arr.reshape(1, -1)).flatten()
        preds_p10.append(inv(p_np[:, 0]))
        preds_p50.append(inv(p_np[:, 1]))
        preds_p90.append(inv(p_np[:, 2]))
        actual_list.append(actuals[i])
        date_list.append(dates[i])

    bt = pd.DataFrame(date_list, columns=[date_col])
    for j, col in enumerate(target_cols):
        bt[f"{col}_actual"] = [a[j] for a in actual_list]
        bt[f"{col}_p10"]    = [p[j] for p in preds_p10]
        bt[f"{col}_p50"]    = [p[j] for p in preds_p50]
        bt[f"{col}_p90"]    = [p[j] for p in preds_p90]
    return bt

def compute_metrics(actual: np.ndarray, pred: np.ndarray) -> dict:
    mae  = np.mean(np.abs(actual - pred))
    rmse = np.sqrt(np.mean((actual - pred) ** 2))
    mask = actual != 0
    mape = np.mean(np.abs((actual[mask] - pred[mask]) / actual[mask])) * 100 if mask.any() else np.nan
    return {"MAE": mae, "RMSE": rmse, "MAPE (%)": mape}

def add_table_to_docx(doc, df: pd.DataFrame):
    df_str = df.astype(str)
    table = doc.add_table(rows=(df.shape[0] + 1), cols=df.shape[1])
    table.style = 'Table Grid'
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = col
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = df_str.iat[i, j]

def main():
    print(f"Loading model from {DEFAULT_CKPT}...")
    model, ckpt, device = load_model(str(DEFAULT_CKPT))
    date_col = ckpt["date_col"]
    feature_cols = ckpt["feature_cols"]
    
    print(f"Scanning datasets in {DATASETS_DIR}...")
    all_files = []
    if DATASETS_DIR.exists():
        for f in DATASETS_DIR.iterdir():
            if f.is_file() and f.name.endswith(('.csv', '.xlsx', '.xls')):
                all_files.append(f)
                
    if not all_files:
        print("No datasets found in datasets folder!")
        return

    # Parse dates to sort files
    file_info_list = []
    for f in all_files:
        try:
            df = read_df_path(f, date_col)
            if len(df) > 0:
                min_date = df[date_col].min()
                max_date = df[date_col].max()
                file_info_list.append({
                    "path": f,
                    "name": f.name,
                    "min_date": min_date,
                    "max_date": max_date,
                    "df": df
                })
        except Exception as e:
            print(f"Error parsing {f.name}: {e}")

    # Sort by oldest date
    file_info_list.sort(key=lambda x: x["min_date"])
    print(f"Found {len(file_info_list)} valid datasets.")

    doc = Document()
    doc.add_heading("Báo Cáo Lịch Sử Dự Báo Giá Xăng Dầu", 0)
    
    for info in file_info_list:
        print(f"Processing: {info['name']} ({info['min_date'].strftime('%Y-%m-%d')} to {info['max_date'].strftime('%Y-%m-%d')})")
        df = info["df"]
        
        # Merge missing exogenous
        df, auto_merged = merge_missing_exo(df, date_col, feature_cols)
        missing = [c for c in feature_cols if c not in df.columns]
        
        doc.add_heading(f"Dataset: {info['name']}", level=1)
        p = doc.add_paragraph(f"Từ ngày: {info['min_date'].strftime('%d/%m/%Y')} - Đến ngày: {info['max_date'].strftime('%d/%m/%Y')}\n")
        
        if missing:
            p.add_run(f"⚠️ Thiếu các cột để chạy model: {', '.join(missing)}\n").bold = True
            continue
            
        future_results, gate_weights = predict_future(model, ckpt, df, device)
        
        p.add_run(f"Dự báo p50 cho {ckpt['horizon']} ngày tiếp theo:\n").bold = True
        
        pred_p50 = future_results["p50"].copy()
        pred_p50[date_col] = pred_p50[date_col].dt.strftime("%d/%m/%Y")
        
        # Format numeric columns to 2 decimal places
        for c in pred_p50.columns:
            if c != date_col:
                pred_p50[c] = pred_p50[c].round(2)
                
        add_table_to_docx(doc, pred_p50)
        doc.add_paragraph("\n")

        # Backtest & Metrics
        if len(df) > ckpt["seq_len"] + 2:
            p.add_run("📉 Chỉ số So sánh Lịch sử (In-sample Backtest):\n").bold = True
            bt_df = backtest(model, ckpt, df, device, n_samples=200)
            
            metrics_data = []
            for col in ckpt["target_cols"]:
                actual = bt_df[f"{col}_actual"].values
                pred   = bt_df[f"{col}_p50"].values
                m = compute_metrics(actual, pred)
                metrics_data.append({
                    "Target": col,
                    "MAE": f"{m['MAE']:.2f}",
                    "RMSE": f"{m['RMSE']:.2f}",
                    "MAPE (%)": f"{m['MAPE (%)']:.2f}%"
                })
            
            metrics_df = pd.DataFrame(metrics_data)
            add_table_to_docx(doc, metrics_df)
            doc.add_paragraph("\n")

    doc.save(OUTPUT_WORD)
    print(f"\nDone! Report saved to {OUTPUT_WORD}")

if __name__ == "__main__":
    main()
